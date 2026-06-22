# -*- coding: utf-8 -*-
"""Build 보고서.docx (plain, no decorative design). Capture areas = bordered empty box + how-to."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# base font
st = doc.styles['Normal']
st.font.name = 'Malgun Gothic'
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def H(text, lvl=1):
    h = doc.add_heading(text, level=lvl)
    for r in h.runs:
        r.font.name = 'Malgun Gothic'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        r.font.color.rgb = RGBColor(0,0,0)
    return h

def P(text='', bold=False, italic=False):
    p = doc.add_para if False else doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = 'Malgun Gothic'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    return p

def CODE(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(text.split('\n')):
        if i: p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = 'Consolas'; r.font.size = Pt(9)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    return p

REPO = r'C:\k-means-clustering'
def CODEFILE(path, start=None, end=None):
    """Embed a source file (or a sub-range) as a code block. start/end = 1-based line number or a substring marker."""
    try:
        lines = open(path, encoding='utf-8').read().replace('\r','').split('\n')
    except Exception as e:
        CODE('(파일을 읽지 못함: %s)' % path); return
    def idx(m, d):
        if m is None: return d
        if isinstance(m, int): return m-1
        for i, l in enumerate(lines):
            if m in l: return i
        return d
    s = idx(start, 0); e = idx(end, len(lines)-1)
    CODE('\n'.join(lines[s:e+1]))

def TABLE(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]; c.paragraphs[0].add_run(htxt).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    # font fix
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Malgun Gothic'; r.font.size=Pt(10)
                    r.element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')
    doc.add_paragraph()
    return t

def CAPTURE(title, height_cm=5.0):
    """Empty bordered box where the user pastes a screenshot/photo."""
    p = doc.add_paragraph(); r = p.add_run(title); r.bold = True
    r.font.name='Malgun Gothic'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    tr = t.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight'); h.set(qn('w:val'), str(int(height_cm*567))); h.set(qn('w:hRule'),'atLeast')
    trPr.append(h)
    pr = cell.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pr.add_run('(여기에 이미지/사진 삽입)')
    run.italic = True; run.font.color.rgb = RGBColor(0x88,0x88,0x88); run.font.size = Pt(9)
    run.font.name='Malgun Gothic'; run.element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')

def HOW(lines):
    p = doc.add_paragraph(); r = p.add_run('캡쳐 방법'); r.bold = True
    r.font.name='Malgun Gothic'; r.element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')
    for ln in lines:
        b = doc.add_paragraph(style='List Bullet'); rr=b.add_run(ln)
        rr.font.name='Malgun Gothic'; rr.font.size=Pt(10)
        rr.element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')
    doc.add_paragraph()

# ===================== TITLE =====================
H('K-means clustering을 위한 FPU IP 설계 — 보고서', 0)
P('학번/이름: __________ (작성 필요)')
P('플랫폼: RVX kmeans_fpu / FPGA: Arty S7-50 (Spartan-7 xc7s50csga324-1) + Pmod OLED RGB + OLIMEX ARM-USB-TINY-H JTAG')
P('각 [캡쳐 N] 박스에 직접 이미지를 붙여넣고, 그 아래 "캡쳐 방법"대로 진행한다.', italic=True)
P('RTL 시뮬 콘솔/파형 캡쳐 절차는 RVX 사용법(documents 디지털회로및시스템설계 Ch12)에 그림과 함께 정리되어 있다. '
  '요약: 콘솔 결과 = make <app>, 파형 = make <app>.debug_view (QuestaSim) → i_platform 에서 i_test1 우클릭 '
  '→ Add Wave → 신호 우클릭 → Radix → float32 → Zoom Full(f).', italic=True)

# ===================== STEP 1 =====================
H('Step 1 — FPU 모듈 동작 확인 (RTL Simulation)', 1)
H('1.1 목표', 2)
P('제공된 IEEE-754 단정밀도 FPU 모듈(adder/multiplier/divider)의 stb/ack handshake 사용법을 익히고, '
  '사칙연산 제어 로직 fpu_top.v의 최소 동작을 RTL 시뮬레이션으로 검증한다.')
H('1.2 작성한 코드 — fpu_top.v', 2)
for s in [
 '입력: clk, rstnn, var_x, var_y, request_fadd/fsub/fmult/fdiv (8개) / 출력: var_z, done',
 'fpu_adder / fpu_multiplier / fpu_divider 를 각각 1개씩만 인스턴스화 (중복 금지)',
 '뺄셈(fsub)은 별도 모듈 없이 x - y = x + (-y) 로 구현: 피연산자 b의 부호비트(bit31)를 반전하여 adder 사용',
 'input_a/b_stb ↔ input_a/b_ack, output_z_stb ↔ output_z_ack handshake를 빠짐없이 사용',
 '연산 완료 시 done을 1-cycle 펄스로 출력하여 상위(IP_TOP)가 결과 시점을 인지',
]:
    b=doc.add_paragraph(style='List Bullet'); rr=b.add_run(s)
    rr.font.name='Malgun Gothic'; rr.font.size=Pt(10); rr.element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')
H('1.2.1 fpu_top.v 전체 코드', 3)
CODEFILE(REPO + r'\hardware\src\fpu_top.v')
H('1.2.2 testbench (tb_fpu_top.v) 코드', 3)
CODEFILE(REPO + r'\hardware\sim\tb_fpu_top.v')
H('1.3 검증 결과 (로컬 iverilog)', 2)
P('입력 x=14.53(0x41687AE1), y=87.91(0x42AFD1EC) 기준, 4종 연산이 IEEE-754 기대값과 일치:')
TABLE(['연산','결과(float)','hex'],
      [['fadd (x+y)','102.44','0x42CCE148'],
       ['fsub (x-y)','-73.38','0xC292C290'],
       ['fmult (x*y)','1277.33','0x449FAAA2'],
       ['fdiv (x/y)','0.165283','0x3E293FDC']])
CAPTURE('[캡쳐 ①] fpu_top 시뮬레이션 콘솔 결과')
HOW(['VSCode PowerShell 터미널에서: cd C:\\k-means-clustering',
     '.\\hardware\\sim\\run.ps1',
     '→ [PASS] fadd/fsub/fmult/fdiv ... 와 RESULT: ALL TESTS PASSED 출력. 콘솔 전체 캡쳐.',
     '주의: bash hardware/sim/run.sh 는 WSL bash에서 깨짐(iverilog 없음+CRLF). PowerShell run.ps1 사용(또는 Git Bash).'])
CAPTURE('[캡쳐 ①-2] (선택) fpu_top 파형 캡쳐', height_cm=4.0)
HOW(['Step 1은 fpu_top 모듈 단독(standalone) 검증이라 RVX debug_view가 아니라 로컬 QuestaSim으로 파형을 뜬다.',
     'cmd/PowerShell에서: cd C:\\k-means-clustering\\hardware\\sim',
     'vsim -do wave_fpu_top.do',
     '→ QuestaSim GUI에서 var_x/var_y/var_z가 float32 radix로, run+Zoom Full 된 파형이 표시됨. Wave 창 캡쳐.',
     'var_z가 fadd 102.44 → fsub -73.38 → fmult 1277.33 → fdiv 0.1653 순으로 바뀜.',
     '참고: GTKWave는 IEEE-754 float 표시가 기본 지원 안 됨 → 강의처럼 float로 보려면 QuestaSim 사용.'])

# ===================== STEP 2 =====================
H('Step 2 — FPU IP 설계 + fpu_test (RTL Sim → FPGA)', 1)
H('2.1 목표', 2)
P('fpu_top을 APB 슬레이브 IP(IP_TOP.v)에 통합하여 프로세서(ORCA)에 연결하고, '
  '사칙연산 테스트 앱(fpu_test.c, 수정 금지)으로 동작을 검증한다.')
H('2.2 작성한 부분 — IP_TOP.v (빈칸)', 2)
P('레지스터 맵: 0x0=x, 0x4=y, 0x8=z(결과), 0xC W=fadd/R=fsub, 0x10 W=fmult/R=fdiv. 채운 핵심 3가지:')
P('① 결과 반영 — APB write가 아닌 경우, FPU 완료 시 결과를 z에 적재', bold=True)
CODE('else if(fpu_done) var_z <= fpu_z;')
P('② fpu_top 인스턴스 + 시작펄스/Busy FSM — 레벨로 유지되는 request(APB wait 동안)를 1-cycle start 펄스로 변환, '
  '연산 중 busy 유지, 같은 접근에서 재트리거 방지(op_serviced).', bold=True)
P('③ rpready_set wait-state — 연산 주소(0xC/0x10) 접근 중 결과 전까지 rpready=0으로 wait-state, 완료 시 1.', bold=True)
CODE('if(rpsel && op_addr && !op_serviced) rpready_set <= 0; // 연산중 -> wait\n'
     'else                                 rpready_set <= 1; // 정상/완료')
P('아래는 IP_TOP.v에서 작성한(빈칸을 채운) 부분이다. (0xC/0x10 연산 요청 디코드는 스켈레톤의 case문에서 처리됨)', italic=True)
H('① FPU 연동 신호 선언', 3); CODEFILE(REPO + r'\hardware\src\IP_TOP.v', 47, 60)
H('② var_z에 FPU 결과 반영 (fpu_done 시 적재)', 3); CODEFILE(REPO + r'\hardware\src\IP_TOP.v', 135, 143)
H('③ start-pulse/busy FSM + fpu_top 인스턴스', 3); CODEFILE(REPO + r'\hardware\src\IP_TOP.v', 147, 196)
H('④ rpready wait-state', 3); CODEFILE(REPO + r'\hardware\src\IP_TOP.v', 200, 207)
H('2.3 SoC 연결 (글루코드)', 2)
P('RVX make syn이 생성한 user/rtl/include/kmeans_fpu_user_region.vh의 주석 템플릿을 IP_TOP(i_test1) 인스턴스로 '
  '연결하여 uNoC(APB)와 special IP를 결선. i_test1 슬레이브 base = 0xE2020000.')
H('글루코드 — kmeans_fpu_user_region.vh (i_test1 = IP_TOP 인스턴스)', 3)
CODEFILE(REPO + r'\rvx_generated\kmeans_fpu\user\rtl\include\kmeans_fpu_user_region.vh', 'IP_TOP', ');')
H('2.4 검증 결과 — RTL Simulation', 2)
P('입력 x=14.53, y=87.91 → fadd 102.44 / fsub -73.38 / fmult 1277.33 / fdiv 0.17(%.2f). Errors 0.')
CAPTURE('[캡쳐 ②] fpu_test RTL 시뮬레이션 콘솔 ([SIM@RTL])')
HOW(['cmd에서: cd C:\\rvx_lec_hw\\platform\\kmeans_fpu\\sim_rtl',
     'make fpu_test',
     '→ 출력 끝의 [SIM@RTL] 블록(x : 14.53 ... fdiv : 0.17)을 캡쳐.'])
CAPTURE('[캡쳐 ②-2] fpu_test 파형 (QuestaSim, RVX debug_view)')
HOW(['RVX 사용법(Ch12)에 나온 절차. cmd에서: cd C:\\rvx_lec_hw\\platform\\kmeans_fpu\\sim_rtl',
     'make fpu_test.debug_view  (QuestaSim GUI가 뜸, 수 분 소요)',
     'i_platform + 펼치기 → i_test1 우클릭 → Add Wave → var_x/var_y/var_z(또는 rprdata) 선택',
     '선택 신호 우클릭 → Radix → float32 → Zoom Full(f) 후 파형 캡쳐 (값 14.53/87.91/102.44…이면 정상)',
     '※ debug_view 중 일시 오류 시 명령을 2~3회 재입력하면 진행됨(Ch12 안내).'])
H('2.5 검증 결과 — FPGA (Arty S7-50)', 2)
P('보드 실측 UART(COM5, 115200). 저장본: rvx_generated/kmeans_fpu/fpga_uart/fpu_test.uart.log')
CODE('[EMU@FPGA]\nx : 14.53\ny : 87.91\nz : 670.72\nfadd : 102.44\nfsub : -73.38\nfmult : 1277.33\nfdiv : 0.17')
CAPTURE('[캡쳐 ③] fpu_test FPGA 실행결과 (PuTTY / UART)')
HOW(['비트스트림 플래시: cd C:\\rvx_lec_hw\\platform\\kmeans_fpu\\imp_arty-50_2026-06-21 후 make program',
     'PuTTY 실행 → Serial / COM5 / 115200 연결 (또는 make printf)',
     '다른 cmd에서 make fpu_test.run 실행 → PuTTY에 위 결과가 뜨면 캡쳐',
     '(PuTTY 대신 fpu_test.uart.log 텍스트를 붙여도 됨)'])
CAPTURE('[캡쳐 ④] FPGA 보드 실물 사진 (본인 이름·학번 함께)', height_cm=6.0)
HOW(['동작 중인 Arty S7-50 보드와 본인 이름·학번 메모(또는 PuTTY 화면)가 한 프레임에 나오도록 사진 1장 촬영.',
     '(Step 3 [캡쳐 ⑧]과 동일 사진 사용 가능)'])

# ===================== STEP 3 =====================
H('Step 3 — K-means clustering에 FPU IP 적용 (RTL Sim → FPGA)', 1)
H('3.1 목표', 2)
P('K-means 응용의 소프트웨어 실수 연산을 설계한 FPU IP의 API로 대체하고, '
  '(1) 군집 결과가 소프트웨어 버전과 동일한지, (2) 가속 효과(total tick)를 확인한다.')
H('3.2 FPU API (fpu_test.c 패턴 재사용)', 2)
P('k_means_oled.c 상단에 헬퍼 추가: set_x/set_y/get_z, perform_fadd/fsub/fmult/fdiv, '
  '그리고 한 번에 호출하는 fpu_add/fpu_sub/fpu_mult/fpu_div(a,b). '
  '(write 0xC=fadd, read 0xC=fsub, write 0x10=fmult, read 0x10=fdiv; 결과는 read 0x8=get_z)')
H('3.3 대체한 실수 연산 4곳 (사칙연산 4종 모두 사용)', 2)
for s in [
 '거리 계산 dis += sqr((float)data[i][l]-means[j][l]) → diff=fpu_sub, sq=fpu_mult(diff,diff), dis=fpu_add(dis,sq)  [fsub·fmult·fadd]',
 '새 mean 누적 temp[group[i]][j] += (float)data[i][j] → fpu_add  [fadd]',
 '평균 temp[i][j] /= count[i] → fpu_div(temp[i][j], (float)count[i])  [fdiv]',
 '수렴 판정 ABS(temp[i][j]-means[i][j])의 뺄셈 → ABS(fpu_sub(temp[i][j], means[i][j]))  [fsub]',
]:
    b=doc.add_paragraph(style='List Bullet'); rr=b.add_run(s)
    rr.font.name='Malgun Gothic'; rr.font.size=Pt(10); rr.element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')
P('또한 profiling_start/end("K-means clustering") 구간 뒤에 profiling_print()를 추가해 clustering 루프의 total tick을 출력.')
H('3.3.1 FPU API 정의 (k_means_oled.c 추가분)', 3)
CODEFILE(REPO + r'\software\k_means_oled.c', '/* ---- FPU IP (i_test1', 'static inline float fpu_div')
H('3.3.2 API로 대체한 실수 연산 4곳', 3)
CODE('// (1) 거리 계산 sqr((float)data - means) : fsub -> fmult -> fadd\n'
     'for(int l = 0; l < 2; l++)\n{\n'
     '    float diff = fpu_sub((float)data[i][l], means[j][l]); // fsub\n'
     '    float sq   = fpu_mult(diff, diff);                    // fmult\n'
     '    dis        = fpu_add(dis, sq);                        // fadd\n}')
CODE('// (2) 새 mean 누적 temp += data : fadd\n'
     'temp[group[i]][j] = fpu_add(temp[group[i]][j], (float)data[i][j]);')
CODE('// (3) 평균 temp /= count : fdiv\n'
     'temp[i][j] = fpu_div(temp[i][j], (float)count[i]);')
CODE('// (4) 수렴 판정 ABS(temp - means) 의 뺄셈 : fsub\n'
     'if(ABS(fpu_sub(temp[i][j], means[i][j])) > 0.0001)')
H('3.4 결과 (1) — 군집 결과가 소프트웨어와 동일 (RTL Sim, num_data=5)', 2)
TABLE(['항목','소프트웨어(float)','FPU IP 적용'],
      [['means[0]','(76.50, 16.00)','(76.50, 16.00)'],
       ['means[1]','(48.00, 39.00)','(48.00, 39.00)'],
       ['means[2]','(50.00, 49.50)','(50.00, 49.50)'],
       ['iteration','2','2'],
       ['group 수','2 / 1 / 2','2 / 1 / 2']])
P('→ 완전히 동일 = IP가 올바르게 설계됨.')
CAPTURE('[캡쳐 ⑤] RTL 시뮬 — 군집 결과 비교 (FPU vs 소프트웨어)')
HOW(['두 앱 src/main.c에서 #define full_printf 1, #define num_data 5 로 둔 뒤 cmd에서:',
     'cd C:\\rvx_lec_hw\\platform\\kmeans_fpu\\sim_rtl',
     'make k_means_oled   (FPU 버전)   /   make k_means_base   (소프트웨어 float 버전)',
     '→ 각 출력의 means / iteration / # of group 이 동일한지 캡쳐.'])
H('3.5 결과 (2) — 가속 효과 (RTL Sim, num_data=5, full_printf=0)', 2)
P('clustering 루프의 profiling tick (루프 내 printf 제외, 순수 연산):')
TABLE(['버전','total tick','time(ms)'],
      [['소프트웨어 float','7057','7.06'],['FPU IP 적용','2880','2.88']])
P('→ 약 2.45배 가속 (가속 정도 자체는 평가 대상 아님).')
CAPTURE('[캡쳐 ⑥] RTL 시뮬 — tick 비교 ([section] K-means clustering total tick)')
HOW(['두 앱 src/main.c에서 #define full_printf 0 으로 바꾼 뒤 위와 동일하게 make k_means_oled / make k_means_base 실행',
     '→ 각 출력의 [section] K-means clustering / total tick (2880, 7057)을 캡쳐.'])
CAPTURE('[캡쳐 ⑥-2] (선택) k_means_oled 파형 (QuestaSim, RVX debug_view)', height_cm=4.0)
HOW(['[캡쳐 ②-2]와 동일 절차로 make k_means_oled.debug_view 실행 후',
     'i_test1의 var_x/var_y/var_z를 Add Wave → Radix float32 → Zoom Full(f)로 캡쳐.',
     '(거리계산/평균에서 FPU 연산이 연속 호출되는 모습 확인 — 없어도 무방)'])
H('3.6 결과 (3) — FPGA 실측 (Arty S7-50, num_data=100)', 2)
P('실데이터 100점에 대해 FPU IP로 군집화가 정상 동작, 3 iterations만에 수렴. '
  '저장본: rvx_generated/kmeans_fpu/fpga_uart/k_means_oled.uart.log')
CODE('[EMU@FPGA]\niteration = 3\nmeans[0] = (70.43, 18.23)\nmeans[1] = (19.72, 35.90)\n'
     'means[2] = (54.08, 49.14)\nEnd of K-means clustering! iteration : 3\n'
     '# of group 0 data : 35\n# of group 1 data : 29\n# of group 2 data : 36\n'
     '[section] K-means clustering  total tick 98389')
CAPTURE('[캡쳐 ⑦] k_means_oled FPGA 실행결과 (PuTTY / UART)')
HOW(['(Step 2에서 플래시 안 했으면) imp 디렉터리에서 make program',
     'PuTTY를 COM5/115200으로 연결 (또는 make printf)',
     'make k_means_oled.run 실행 → PuTTY에 위 군집 결과가 뜨면 캡쳐',
     '(PuTTY 대신 k_means_oled.uart.log 텍스트로 대체 가능)'])
CAPTURE('[캡쳐 ⑧] FPGA 보드 실물 사진 (본인 이름·학번 함께)', height_cm=6.0)
HOW(['[캡쳐 ④]와 동일 사진 사용 가능.'])

# ===================== CHECKLIST =====================
H('제출물 체크리스트', 1)
for s in [
 '코드: submission/step1_code.zip, step2_code.zip, step3_code.zip (제출 전 [학번이름]stepN_코드.zip 으로 변경)',
 '보고서: 본 문서에 [캡쳐 ①~⑧] 채워 제출',
 '필수 직접 캡쳐: ① ② ⑤ ⑥ (콘솔), ③ ⑦ (UART), ④/⑧ (보드+이름·학번 사진)',
]:
    b=doc.add_paragraph(style='List Bullet'); rr=b.add_run(s)
    rr.font.name='Malgun Gothic'; rr.font.size=Pt(10); rr.element.rPr.rFonts.set(qn('w:eastAsia'),'Malgun Gothic')

import sys
out = sys.argv[1] if len(sys.argv) > 1 else r'C:\k-means-clustering\submission\보고서.docx'
doc.save(out)
print('saved:', out)
